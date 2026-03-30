"""
EcoGuard - Documentation Generator
Produces a comprehensive Word (.docx) file covering every aspect of the system.
Run: python generate_docs.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ────────────────────────────────────────────────
# Helpers
# ────────────────────────────────────────────────
def set_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h

def add_para(doc, text, bold=False, italic=False, size=None, color=None, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    return p

def add_code(doc, text):
    """Add a monospace code block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x44)
    # Light grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p

def add_table_row(table, cells, bold_first=False):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = str(text)
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        if bold_first and i == 0:
            cell.paragraphs[0].runs[0].bold = True
    return row

def shade_row(row, hex_color='D9EAD3'):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

# ────────────────────────────────────────────────
# Document
# ────────────────────────────────────────────────
doc = Document()

# Page margins
sections = doc.sections
for section in sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── COVER PAGE ──────────────────────────────────
doc.add_picture  # no image, just text cover
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("EcoGuard Technologies")
r.font.size = Pt(28)
r.bold = True
r.font.color.rgb = RGBColor(0x00, 0x55, 0x88)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("Pipeline Integrity Platform")
r2.font.size = Pt(18)
r2.font.color.rgb = RGBColor(0x33, 0x77, 0xAA)

doc.add_paragraph()
t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run("COMPLETE TECHNICAL DOCUMENTATION")
r3.font.size = Pt(13)
r3.bold = True
r3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = t4.add_run(f"Version 1.0  ·  {datetime.date.today().strftime('%B %d, %Y')}")
r4.font.size = Pt(11)
r4.italic = True

doc.add_page_break()

# ── SECTION 1: WHAT IS ECOGUARD ─────────────────
set_heading(doc, "1.  WHAT IS ECOGUARD?", 1, (0x00, 0x55, 0x88))

add_para(doc,
    "EcoGuard is an industrial-grade pipeline leakage detection and monitoring system. "
    "It simulates a real 48 km oil/gas pipeline network monitored by 4 intelligent sensors. "
    "The system uses autonomous AI agents, real-time WebSocket streaming, and a professional "
    "mobile dashboard to detect, alert, and respond to pipeline leaks before they cause "
    "environmental damage.", size=11)

doc.add_paragraph()
add_para(doc, "Why It Exists", bold=True, size=11)
add_para(doc,
    "Pipeline leaks cost billions in cleanup costs and environmental fines every year. "
    "Most leaks go undetected for hours because monitoring systems generate too many false alarms. "
    "EcoGuard combines 4 types of sensors with weighted risk scoring to give "
    "accurate, real-time leak probability — reducing false alarms while catching real events fast.", size=10)

doc.add_paragraph()
add_para(doc, "Key Technical Achievements", bold=True, size=11)
for item in [
    "4 autonomous sensor agents running as async background tasks in Python",
    "Real-time WebSocket data streaming to mobile app (2-second tick interval)",
    "Z-score statistical anomaly detection per sensor",
    "Composite weighted risk fusion: 35% pressure + 30% flow + 25% acoustic + 10% infrared",
    "Full alert workflow: create → WebSocket broadcast → acknowledge → respond → resolve",
    "Dark / Light / System theme support with global React Context",
    "React Native (Expo SDK 54) mobile app — runs on iPhone XR via Expo Go",
    "FastAPI async backend with SQLite persistence (PostgreSQL-ready)",
    "6-screen professional industrial UI with live visualizations",
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item).font.size = Pt(10)

doc.add_page_break()

# ── SECTION 2: SYSTEM ARCHITECTURE ──────────────
set_heading(doc, "2.  SYSTEM ARCHITECTURE", 1, (0x00, 0x55, 0x88))

add_para(doc, "How the two halves talk to each other:", bold=True, size=11)

add_code(doc,
"iPhone XR (Expo Go app)          School WiFi           Laptop (Backend)\n"
"       |                             |                        |\n"
"   Mobile App  ←── HTTP REST ────────────────→  FastAPI Server :8000\n"
"   (React Native)  ←── WebSocket ─────────────→  4x Sensor Agents\n"
"                                                  SQLite Database")

doc.add_paragraph()
add_para(doc, "Communication Protocol:", bold=True, size=11)

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
hdr = t.rows[0].cells
hdr[0].text = "Protocol"
hdr[1].text = "Used For"
hdr[2].text = "Address"
shade_row(t.rows[0], 'BDD7EE')
for row_data in [
    ("HTTP REST", "API calls (fetch sensors, create reports, etc.)", "http://IP:8000/api/..."),
    ("WebSocket", "Live streaming every 2 seconds + instant alerts", "ws://IP:8000/ws/pipeline"),
    ("HTTP GET", "Health check endpoint", "http://IP:8000/health"),
]:
    add_table_row(t, row_data)

doc.add_paragraph()
set_heading(doc, "2.1  Frontend: React Native Mobile App", 2, (0x00, 0x77, 0xAA))
add_para(doc,
    "Located at: mobile/  ·  Framework: Expo SDK 54  ·  Language: TypeScript  ·  Navigation: Expo Router v4", size=10, italic=True)
add_para(doc,
    "The mobile app is built with React Native (Expo). It runs natively on iOS via Expo Go and also "
    "in any desktop web browser. The app has 6 main screens organized as tabs at the bottom. "
    "State is managed through React Context providers that wrap the entire app.", size=10)

doc.add_paragraph()
set_heading(doc, "2.2  Backend: FastAPI Server", 2, (0x00, 0x77, 0xAA))
add_para(doc,
    "Located at: backend/  ·  Framework: FastAPI (Python)  ·  Database: SQLite  ·  ORM: SQLAlchemy", size=10, italic=True)
add_para(doc,
    "The backend is an async Python server. On startup, it seeds the database with 4 sensors and "
    "9 pipeline nodes, then launches 4 autonomous sensor agents as asyncio background tasks. "
    "These agents generate realistic sensor readings every 2 seconds and broadcast them via WebSocket.", size=10)

doc.add_page_break()

# ── SECTION 3: THE 4 SENSORS ─────────────────────
set_heading(doc, "3.  THE 4 SENSORS — What They Are and What They Detect", 1, (0x00, 0x55, 0x88))

add_para(doc,
    "In a real pipeline, these are physical hardware devices installed at key points along the pipe. "
    "In EcoGuard, they are simulated by software agents that generate statistically realistic data "
    "with Gaussian noise (random variation) and sinusoidal drift (slow natural fluctuation).", size=10)

doc.add_paragraph()

sensors = [
    ("PS-001 — Pressure Transducer",
     "Km 12.4", "PSI (Pounds per Square Inch)", "85 – 115 PSI",
     "Cyan (#00E5FF)",
     "Measures how hard the fluid is pushing against the pipe walls. "
     "When a pipe cracks or a valve fails, pressure DROPS because fluid escapes. "
     "PS-001 drops 2.5 PSI per tick (every 2 seconds) during a leak, reaching as low as ~76 PSI.",
     "PRESSURE DROP from 103 PSI to 76 PSI = 26% drop in 14 seconds"),
    ("FS-002 — Turbine Flow Meter",
     "Km 22.1", "L/min (Litres per Minute)", "180 – 220 L/min",
     "Amber (#FFAB00)",
     "Measures how much fluid flows through the pipe per minute. "
     "A leak means fluid is escaping before reaching the meter, so flow DECREASES. "
     "During a leak, FS-002 drops to ~70% of nominal (126 L/min from 200 L/min).",
     "FLOW DEVIATION from 200 L/min to 126 L/min = 37% reduction"),
    ("ACS-003 — Acoustic Emission Sensor",
     "Km 35.7", "dB (Decibels)", "28 – 55 dB",
     "Green (#00E676)",
     "Listens for high-frequency sound waves travelling through the pipe wall. "
     "When fluid escapes through a crack at high pressure, it creates a hissing sound (turbulent flow noise). "
     "ACS-003 SPIKES +35 dB on a leak event, reaching ~90 dB (louder than a lawn mower).",
     "ACOUSTIC SPIKE from 38 dB to 73 dB = +93% increase"),
    ("IR-004 — Infrared Thermographic Array",
     "Km 44.2", "°C (Celsius)", "15 – 28°C",
     "Magenta (#E040FB)",
     "Scans the ground surface temperature around the pipeline. "
     "When pressurized fluid escapes into the surrounding soil, friction and chemical reaction "
     "create heat zones (thermal anomalies). IR-004 RAMPS +0.6°C per tick, reaching +18°C above baseline.",
     "THERMAL RAMP from 22°C to 40°C = +82% elevation over 60 seconds"),
]

for sname, loc, unit, nominal, color, desc, example in sensors:
    set_heading(doc, f"3.x  {sname}", 2, (0x00, 0x66, 0x99))
    t = doc.add_table(rows=5, cols=2)
    t.style = 'Table Grid'
    labels = ["Location", "Measurement Unit", "Normal Range (Nominal)", "UI Accent Color", "Leak Signature"]
    values = [loc, unit, nominal, color, example]
    for i, (lbl, val) in enumerate(zip(labels, values)):
        row = t.rows[i]
        row.cells[0].text = lbl
        row.cells[1].text = val
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        if i == 0:
            shade_row(row, 'E8F4FD')
    add_para(doc, desc, size=10)
    doc.add_paragraph()

doc.add_page_break()

# ── SECTION 4: LEAK DETECTION ALGORITHM ──────────
set_heading(doc, "4.  HOW LEAK DETECTION WORKS — The Algorithm", 1, (0x00, 0x55, 0x88))

set_heading(doc, "4.1  Step 1 — Z-Score Anomaly Detection (per sensor)", 2, (0x00, 0x77, 0xAA))
add_para(doc,
    "Z-Score measures how far a reading is from the sensor's historical average, in units of standard deviation. "
    "A Z-Score above 3.0 (or below -3.0) means the reading is statistically unusual — flagged as an anomaly.", size=10)

add_code(doc,
"Formula:   Z = (current_value - rolling_mean) / rolling_std\n"
"\n"
"Example:   PS-001 mean = 100 PSI,  std = 4 PSI\n"
"           Current reading = 76 PSI\n"
"           Z = (76 - 100) / 4 = -6.0   ← Well above threshold of 3.0\n"
"           Result: is_anomaly = True,  confidence_score = 0.97")

doc.add_paragraph()
set_heading(doc, "4.2  Step 2 — Leak Probability per Sensor (0.0 to 1.0)", 2, (0x00, 0x77, 0xAA))
add_para(doc,
    "Each sensor computes its own leak probability based on: (1) whether it's in LEAK agent state, "
    "(2) the severity of the anomaly (Z-score magnitude), and (3) the sensor type's characteristic "
    "response to leaks. Values are stored as decimals: 0.0 = no leak, 1.0 = definite leak.", size=10)

doc.add_paragraph()
set_heading(doc, "4.3  Step 3 — Composite Risk Fusion (all 4 sensors combined)", 2, (0x00, 0x77, 0xAA))
add_para(doc,
    "Real leak detection requires agreement from MULTIPLE sensors. A single sensor anomaly could be "
    "an equipment glitch. Four sensors all showing anomalies = high confidence of a real leak.", size=10)

add_code(doc,
"Composite Risk = (P_prob × 0.35) + (F_prob × 0.30) + (A_prob × 0.25) + (IR_prob × 0.10)\n"
"\n"
"Where:\n"
"  P_prob  = Pressure sensor leak probability (weight: 35% — most reliable indicator)\n"
"  F_prob  = Flow sensor leak probability     (weight: 30% — second most reliable)\n"
"  A_prob  = Acoustic sensor leak probability (weight: 25% — fast detection)\n"
"  IR_prob = Infrared sensor leak probability (weight: 10% — delayed response)\n"
"\n"
"Risk Levels:\n"
"  0 – 30%  → NOMINAL (green)\n"
"  30 – 70% → WARNING (amber) — investigate\n"
"  70 – 100%→ CRITICAL (red)  — immediate response required")

doc.add_paragraph()
add_para(doc,
    "Why these weights? Pressure drops first and most reliably when a pipe breaches. "
    "Flow rate confirms the material loss. Acoustic emission detects the turbulent hissing sound quickly. "
    "Infrared takes longer to respond as soil heats up, so it is weighted less.", size=10)

doc.add_paragraph()
set_heading(doc, "4.4  Step 4 — Alert Generation with Debounce", 2, (0x00, 0x77, 0xAA))
add_para(doc,
    "When an anomaly is detected, the alert service creates an alert in the database "
    "AND broadcasts it via WebSocket. To prevent alert fatigue (too many duplicate alerts), "
    "a 30-second debounce timer prevents the same sensor from generating a new alert "
    "until the previous one has been active for at least 30 seconds.", size=10)

doc.add_page_break()

# ── SECTION 5: MOBILE APP SCREENS ────────────────
set_heading(doc, "5.  MOBILE APP — All 6 Screens Explained", 1, (0x00, 0x55, 0x88))

screens = [
    ("5.1  Dashboard (Home Screen)", (0x00, 0x77, 0xAA),
     "At-a-glance status of the entire pipeline system. This is the first screen you see.",
     [
         ("ECOGUARD header", "Top of screen. Shows 'ECOGUARD' title in cyan, LIVE/OFFLINE badge, and theme toggle button."),
         ("SIMULATE LEAK button", "Red button. Injects a fake leak event into ALL 4 sensors simultaneously for ~30 seconds. Used for testing and demonstration."),
         ("RESET SYSTEM button", "Grey button. Cancels any active leak injection and returns all agents to nominal state."),
         ("KPI Cards (4 boxes)", "KPI = Key Performance Indicator. Shows CURRENT live value for each sensor: Pressure (PSI), Flow Rate (L/min), Acoustic (dB), Infrared (°C). Updates every 2 seconds. Live dot pulses when connected."),
         ("Composite Leak Risk Gauge", "The big animated arc (semicircle). Shows overall pipeline risk as a percentage 0-100%. Needle sweeps right as risk increases. Green (low) → Amber (warning) → Red (critical). Label below shows NOMINAL / WARNING / CRITICAL."),
         ("Sensor Array — Real-Time", "4 mini cards showing all sensors at once. Each card shows current value, a small sparkline bar chart of last 20 readings, and leak probability text."),
         ("Live Sensor Trend Chart", "Line chart plotting all 4 sensor values over recent history. Each sensor has its own color line. Legend at bottom identifies each line."),
         ("Event Log", "List of the 10 most recent alerts. Shows time, sensor, alert type, and value. Red = critical, amber = warning."),
     ]),
    ("5.2  Analytics Screen", (0xAA, 0x00, 0xCC),
     "Deep historical analysis of each sensor. Used to understand past behaviour and investigate what caused alerts.",
     [
         ("DATA WINDOW selector (1H/6H/24H/48H)", "Choose how far back to look: 1 hour, 6 hours, 24 hours, or 48 hours. Tap a button to change the time window. The charts and statistics update accordingly."),
         ("Per-Sensor Cards (4 cards)", "One card per sensor. Each shows statistics and a chart."),
         ("Statistics row (MEAN/STD DEV/MIN/MAX/READINGS)", "MEAN = average value over the selected time window. STD DEV = standard deviation (how much values varied). MIN/MAX = lowest and highest values seen. READINGS = total number of data points collected."),
         ("Line Chart", "Plots sensor values over time. X-axis = time, Y-axis = sensor value. Each point on the line = one 2-second reading."),
         ("Red dots on chart", "ANOMALY MARKERS. Each red dot = a reading that was flagged by Z-score detection as an anomaly. If you see a cluster of red dots, that is when the leak event was active."),
         ("ANOMALIES DETECTED banner", "Red bar at bottom of card showing how many anomalies were detected in the selected window."),
     ]),
    ("5.3  Alerts Screen", (0xCC, 0x00, 0x00),
     "Real-time incident management. New alerts appear here instantly via WebSocket as soon as the backend detects an anomaly.",
     [
         ("Stats Row (CRITICAL / WARNINGS / UNACKED)", "CRITICAL = count of highest severity active alerts. WARNINGS = count of medium severity. UNACKED = count of alerts you have not yet acknowledged."),
         ("Filter buttons (ALL / CRITICAL / WARNING / HISTORY)", "Filter which alerts to show. HISTORY button switches from 'active only' to 'all alerts ever'. Shows up to 25 per page with 'Load More'."),
         ("Alert Cards", "Each card = one alert event. Left colored bar = severity. Title = alert type (e.g. ABNORMAL PRESSURE DROP). Description = detailed message from the system. Meta row = sensor ID, time ago, and the value that triggered the alert."),
         ("Unread dot", "Red dot next to the timestamp means the alert has NOT been acknowledged yet."),
         ("Expand card (tap)", "Tap any alert to expand it and see the action panel."),
         ("ACKNOWLEDGE button", "Marks the alert as seen. Turns off the red unread dot. Does NOT resolve the incident — just confirms you are aware."),
         ("RESPONSE ACTIONS (4 buttons)", "ISOLATE VALVE = close the nearest isolation valve to stop flow. DISPATCH TEAM = send a field crew to inspect. REDUCE PRESSURE = lower system pressure to prevent worsening. NOTIFY HQ = send notification to headquarters. Tapping confirms and resolves the alert."),
         ("LOAD MORE button", "Appears when there are more than 25 alerts. Tap to load the next 25."),
     ]),
    ("5.4  Pipeline Map Screen", (0x00, 0x99, 0x88),
     "Visual topology of the 48 km pipeline from source pump to receiving terminal. Shows WHERE a leak is on the pipeline.",
     [
         ("Stats Row (48km / 9 Nodes / Leak Risk %)", "Quick reference numbers. 48km = total pipeline length. 9 = number of monitoring nodes. Risk % = composite leak risk from the detection algorithm."),
         ("Pipeline Topology SVG Map", "A technical drawing of the pipeline. Circles = nodes (monitoring points). Lines connecting them = pipe segments. Color of lines and circles changes based on leak risk at each point."),
         ("Node colors", "GREEN = normal operation. AMBER = warning level anomaly detected. RED = critical anomaly / high leak probability. GREY = isolated (valve closed)."),
         ("Node symbols", "P = Pump (source of flow). J = Junction (pipe splits). V = Valve (flow control). S = Sensor Point (where a sensor is installed). E = Endpoint (receiving terminal)."),
         ("Km markers", "Small numbers under nodes showing distance from source in kilometres."),
         ("Tap a node", "Tap any circle on the map to open a detail panel at the bottom. Shows node type, status, km marker. If a sensor is attached, also shows current reading, leak probability, and anomaly flag."),
         ("Map Legend", "Key explaining what each color and symbol means."),
         ("Sensor List", "Below the legend: lists all 4 sensors with their location in the pipeline."),
     ]),
    ("5.5  Sensors Screen", (0x00, 0xBB, 0x55),
     "Detailed technical view of each sensor with its own live visualization. The most detailed view of individual sensor data.",
     [
         ("PS-001 — Pressure Gauge (semicircle)", "Animated arc gauge. The needle rotates right as pressure increases. The colored arc fills as pressure approaches maximum. Shows current PSI value below."),
         ("FS-002 — Flow Rate Bars", "20 vertical bars. Each bar = one reading. Height = flow rate. Bars are updated in real-time, scrolling left as new readings arrive. Shows current L/min."),
         ("ACS-003 — Acoustic Waveform", "Oscilloscope-style waveform line. Shows the acoustic signal shape over the last 32 readings. During a leak, you see a spike in the waveform. Shows current dB value."),
         ("IR-004 — Infrared Heat Map", "24 colored bars. Color changes from TEAL (cool/normal) → AMBER (warm/warning) → RED (hot/critical) as temperature rises. Shows current °C."),
         ("Spec grid (4 items)", "LOCATION = where on the 48km pipeline this sensor is installed. NOMINAL RANGE = the expected normal operating range. AGENT STATE = NORMAL (no leak) or LEAK (leak mode active). CONFIDENCE = how confident the detection algorithm is in its assessment (0–100%)."),
         ("LEAK PROBABILITY bar", "Horizontal progress bar. Green when below 40%, amber between 40-70%, red above 70%. Percentage shown on the right."),
         ("Description text", "Technical specification of the sensor hardware type and detection principle."),
     ]),
    ("5.6  Reports Screen", (0xCC, 0x88, 0x00),
     "Generate and store compliance reports. Used for regulatory requirements, incident documentation, and health monitoring.",
     [
         ("Sensor Health — 24H", "Table showing each sensor's uptime and health score over the last 24 hours. The bar fills green/amber/red based on anomaly-free reading percentage. READINGS column shows total data points collected."),
         ("Generate Report buttons (5 types)", "DAILY SUMMARY = 24-hour overview of all sensor activity. WEEKLY REPORT = 7-day analysis. INCIDENT REPORT = focused on alerts and anomalies in last 24 hours. COMPLIANCE EXPORT = 30-day regulatory compliance summary. SENSOR HEALTH = detailed health metrics. Tap to generate. Shows 'GENERATING...' while working."),
         ("Generated Reports list", "Table of all reports stored in the database. Shows type, title, and generation timestamp. Each row has a JSON badge — the report is stored as a JSON file."),
     ]),
]

for title, color, intro, items in screens:
    set_heading(doc, title, 2, color)
    add_para(doc, intro, size=10, italic=True)
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    t.rows[0].cells[0].text = "UI Element"
    t.rows[0].cells[1].text = "What It Means and Why It Is There"
    shade_row(t.rows[0], 'BDD7EE')
    for i in range(2):
        t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        t.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    for elem, meaning in items:
        row = t.add_row()
        row.cells[0].text = elem
        row.cells[1].text = meaning
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()

doc.add_page_break()

# ── SECTION 6: HEADER / THEME TOGGLE ─────────────
set_heading(doc, "6.  HEADER ELEMENTS EXPLAINED", 1, (0x00, 0x55, 0x88))

items_header = [
    ("ECOGUARD title (top-left)", "The app title. On the Dashboard it says 'ECOGUARD'. On other tabs it says the screen name (ANALYTICS, ALERTS, etc.)."),
    ("LIVE / OFFLINE badge", "Shows WebSocket connection status. GREEN dot + 'LIVE' = receiving real-time data from backend. RED dot + 'OFFLINE' = connection lost. 'RECONNECTING' = trying to reconnect (happens if backend restarts)."),
    ("DARK / LIGHT / SYSTEM button (moon icon)", "Theme toggle. Tap to cycle between: DARK (dark navy/cyan industrial theme), LIGHT (white/blue clean theme), SYSTEM (follows your iPhone's light/dark mode setting). The icon changes: moon = dark, sun = light, phone = system. The ENTIRE app instantly changes theme — not just the header."),
]

t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
t.rows[0].cells[0].text = "Header Element"
t.rows[0].cells[1].text = "Explanation"
shade_row(t.rows[0], 'BDD7EE')
for i in range(2):
    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    t.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(9)
for elem, meaning in items_header:
    row = t.add_row()
    row.cells[0].text = elem
    row.cells[1].text = meaning
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)

doc.add_page_break()

# ── SECTION 7: HOW TO RUN ─────────────────────────
set_heading(doc, "7.  HOW TO RUN ECOGUARD", 1, (0x00, 0x55, 0x88))

set_heading(doc, "7.1  Prerequisites", 2, (0x00, 0x77, 0xAA))
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
t.rows[0].cells[0].text = "Requirement"
t.rows[0].cells[1].text = "Version"
t.rows[0].cells[2].text = "Where to Get"
shade_row(t.rows[0], 'BDD7EE')
for i in range(3):
    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    t.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(9)
for row_data in [
    ("Python", "3.10+", "python.org"),
    ("Node.js", "18+", "nodejs.org"),
    ("Expo Go (iPhone)", "SDK 54", "App Store → search 'Expo Go'"),
    ("Both devices on same WiFi", "—", "School WiFi: 'Student'"),
]:
    row = t.add_row()
    for i, text in enumerate(row_data):
        row.cells[i].text = text
        row.cells[i].paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph()
set_heading(doc, "7.2  One-Time Setup", 2, (0x00, 0x77, 0xAA))
add_para(doc, "Step 1: Get your WiFi IP address", bold=True, size=10)
add_code(doc, "Run in Command Prompt:\n  ipconfig\n\nLook for: Wireless LAN adapter WiFi → IPv4 Address\nExample: 10.107.57.10\n\nIMPORTANT: NOT the VMware adapter (192.168.x.x) — that is virtual!")

add_para(doc, "Step 2: Set the IP in the environment file", bold=True, size=10)
add_code(doc, "Edit: mobile/.env.local\n\nChange these two lines:\n  EXPO_PUBLIC_API_URL=http://YOUR_WIFI_IP:8000\n  EXPO_PUBLIC_WS_URL=ws://YOUR_WIFI_IP:8000\n\nExample:\n  EXPO_PUBLIC_API_URL=http://10.107.57.10:8000\n  EXPO_PUBLIC_WS_URL=ws://10.107.57.10:8000")

add_para(doc, "Step 3: Install backend dependencies", bold=True, size=10)
add_code(doc, "cd backend\npip install -r requirements.txt")

add_para(doc, "Step 4: Install frontend dependencies", bold=True, size=10)
add_code(doc, "cd mobile\nnpm install --legacy-peer-deps")

doc.add_paragraph()
set_heading(doc, "7.3  Daily Startup (Every Time)", 2, (0x00, 0x77, 0xAA))
add_para(doc, "Terminal 1 — Start Backend:", bold=True, size=10)
add_code(doc, "cd c:\\Users\\hp\\OneDrive\\Desktop\\projects\\ECOGUARD\\backend\npython main.py\n\nExpected output:\n  INFO: Uvicorn running on http://0.0.0.0:8000\n  INFO: EcoGuard v1.0.0 started — 4 sensor agents active.")

add_para(doc, "Terminal 2 — Start Frontend:", bold=True, size=10)
add_code(doc, "cd c:\\Users\\hp\\OneDrive\\Desktop\\projects\\ECOGUARD\\mobile\nnpm start\n\nExpected output:\n  Expo DevTools running at http://localhost:8081\n  QR code shown below — scan with Expo Go")

add_para(doc, "Terminal 3 (Optional) — Verify Backend:", bold=True, size=10)
add_code(doc, "curl http://10.107.57.10:8000/health\n\nExpected:\n  {\"status\": \"operational\", \"active_agents\": 4}")

add_para(doc, "On iPhone:", bold=True, size=10)
add_para(doc, "1. Open Expo Go app  →  2. Tap Scan QR Code  →  3. Scan the QR from Terminal 2  →  4. Wait 30-60 seconds  →  5. See the live dashboard", size=10, indent=True)

doc.add_page_break()

# ── SECTION 8: SIMULATION ─────────────────────────
set_heading(doc, "8.  HOW TO SIMULATE A LEAK", 1, (0x00, 0x55, 0x88))

add_para(doc, "Experiment 1: Basic Leak Test", bold=True, size=11)
steps_1 = [
    "Open the Dashboard tab on your phone",
    "Note current Leak Risk % on the gauge (should be ~5-15%)",
    "Tap the red SIMULATE LEAK button",
    "Confirm in the popup dialog",
    "Watch the Leak Risk Gauge rise toward 100% over ~10 seconds",
    "Go to the Alerts tab — you should see new CRITICAL alerts appearing",
    "Go to the Analytics tab — check PS-001 chart for red anomaly dots",
    "Go to the Sensors tab — watch the Pressure gauge needle drop",
    "After ~30 seconds, tap RESET SYSTEM on Dashboard",
    "Watch all values return to normal",
]
for i, step in enumerate(steps_1, 1):
    p = doc.add_paragraph(f"{i}.  {step}", style='List Number')
    p.runs[0].font.size = Pt(10)

doc.add_paragraph()
add_para(doc, "Experiment 2: Respond to an Alert", bold=True, size=11)
steps_2 = [
    "Simulate a leak (see Experiment 1)",
    "Go to the Alerts tab",
    "Find a CRITICAL alert — tap it to expand",
    "Tap ACKNOWLEDGE → confirms you are aware",
    "Tap ISOLATE VALVE → confirm in dialog",
    "Alert moves to resolved state",
    "In backend logs (Terminal 1), see: 'resolve' API called",
]
for i, step in enumerate(steps_2, 1):
    p = doc.add_paragraph(f"{i}.  {step}", style='List Number')
    p.runs[0].font.size = Pt(10)

doc.add_paragraph()
add_para(doc, "Experiment 3: Read the Analytics", bold=True, size=11)
steps_3 = [
    "Simulate a leak and wait 30 seconds",
    "Go to Analytics tab",
    "Select '1H' time window",
    "Look at PS-001 (Pressure) — see the dip in the chart",
    "Red dots = anomaly readings during the leak",
    "Check MEAN value — it should be lower than normal (leak pulled it down)",
    "Check STD DEV — higher than normal means more variability (leak caused instability)",
]
for i, step in enumerate(steps_3, 1):
    p = doc.add_paragraph(f"{i}.  {step}", style='List Number')
    p.runs[0].font.size = Pt(10)

doc.add_page_break()

# ── SECTION 9: TROUBLESHOOTING ────────────────────
set_heading(doc, "9.  TROUBLESHOOTING", 1, (0x00, 0x55, 0x88))

problems = [
    ("exp://192.168.232.1:8081 — Timeout error on iPhone",
     "Root Cause", "Expo is using the VMware virtual network adapter instead of your real WiFi. This IP (192.168.x.x from VMware) is not reachable by your iPhone.",
     "Fix",
     "Run ipconfig, find your WiFi IPv4 (e.g. 10.107.57.10).\n"
     "Edit mobile/.env.local — set EXPO_PUBLIC_API_URL=http://10.107.57.10:8000\n"
     "Stop Expo (Ctrl+C), restart: npm start\n"
     "The QR code URL should now show exp://10.107.57.10:8081"),
    ("Project incompatible — SDK 52 vs SDK 54",
     "Root Cause", "Expo Go on your iPhone is SDK 54. The project was configured for SDK 52.",
     "Fix",
     "package.json has been updated to expo ~54.0.0 and all dependent packages.\n"
     "Run: cd mobile && npm install --legacy-peer-deps\n"
     "Then restart: npm start"),
    ("Data shows '--' and Live indicator is red (OFFLINE)",
     "Root Cause", "WebSocket cannot connect to backend. Either backend is not running, or wrong IP.",
     "Fix",
     "1. Make sure backend is running: python main.py\n"
     "2. Check mobile/.env.local has correct WiFi IP\n"
     "3. Verify: curl http://YOUR_IP:8000/health → should return OK\n"
     "4. Check Windows Firewall allows port 8000 (Control Panel → Firewall → Allow app)"),
    ("Backend shows: WARNING: Invalid HTTP request received",
     "Root Cause", "ngrok's HTTPS tunnel sends TLS handshake bytes to the HTTP uvicorn server. These are harmless health-check probes.",
     "Fix",
     "This is now suppressed in main.py by the _SuppressBadRequestFilter.\n"
     "Restart backend: python main.py\n"
     "You will no longer see these warnings. API calls still work fine."),
    ("Theme toggle only changes header/footer, not the whole app",
     "Root Cause", "Screens were using hardcoded COLORS constants instead of useTheme() hook. Now fixed.",
     "Fix",
     "All 5 tab screens (alerts, analytics, sensors, pipeline-map, reports) have been updated.\n"
     "Reload the app. Tapping the theme button now changes the entire app background, text, borders, and accents."),
    ("Reports page shows content duplicated multiple times",
     "Root Cause", "Missing width style on health bars and no deduplication guard on data fetching.",
     "Fix",
     "reports.tsx now includes: (1) width: health+'%' on fill bars, (2) mountedRef guard to prevent double-fetch, (3) array deduplication by id/sensor_id."),
    ("npm install fails with peer dependency errors",
     "Root Cause", "React Native ecosystem has strict peer dependency requirements.",
     "Fix",
     "Always use: npm install --legacy-peer-deps\n"
     "If still failing: rm -rf node_modules package-lock.json && npm cache clean --force && npm install --legacy-peer-deps"),
]

for prob, root_label, root, fix_label, fix in problems:
    set_heading(doc, prob, 3, (0xAA, 0x22, 0x00))
    t = doc.add_table(rows=2, cols=2)
    t.style = 'Table Grid'
    t.rows[0].cells[0].text = root_label
    t.rows[0].cells[1].text = root
    t.rows[1].cells[0].text = fix_label
    t.rows[1].cells[1].text = fix
    shade_row(t.rows[0], 'FCE5CD')
    shade_row(t.rows[1], 'D9EAD3')
    for r in t.rows:
        r.cells[0].paragraphs[0].runs[0].bold = True
        r.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        r.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()

doc.add_page_break()

# ── SECTION 10: BACKEND API REFERENCE ────────────
set_heading(doc, "10.  BACKEND API REFERENCE", 1, (0x00, 0x55, 0x88))
add_para(doc, "All API endpoints are accessible at http://YOUR_IP:8000/docs (Swagger UI when backend is running).", size=10, italic=True)

endpoints = [
    ("GET",  "/health",                           "Health check — returns agent count and connection count"),
    ("GET",  "/api/sensors",                      "List all 4 sensors with their metadata"),
    ("GET",  "/api/sensors/{id}",                 "Get one sensor by ID (e.g. PS-001)"),
    ("POST", "/api/sensors/{id}/inject-leak",     "Inject a leak into one sensor. ?duration_ticks=15"),
    ("POST", "/api/sensors/inject-all-leaks",     "Inject leak into all 4 sensors simultaneously"),
    ("POST", "/api/sensors/reset-all",            "Reset all agents to nominal state"),
    ("GET",  "/api/readings/latest",              "Latest reading for each of the 4 sensors"),
    ("GET",  "/api/readings/kpi",                 "KPI snapshot (current values + status for dashboard)"),
    ("GET",  "/api/readings/{sensor_id}",         "Historical readings for one sensor. ?hours=1&limit=200"),
    ("GET",  "/api/alerts",                       "List alerts. ?active=true for active only, ?hours=48"),
    ("POST", "/api/alerts/{id}/acknowledge",      "Mark alert as acknowledged"),
    ("POST", "/api/alerts/{id}/resolve",          "Resolve alert with action taken. Body: {action_taken}"),
    ("GET",  "/api/analytics/{id}/trend",         "Trend data for charts. ?hours=24"),
    ("GET",  "/api/analytics/{id}/anomalies",     "List of anomalous readings. ?hours=48"),
    ("GET",  "/api/analytics/{id}/statistics",    "Statistical summary: mean, std, min, max, anomaly_count"),
    ("GET",  "/api/analytics/composite/risk",     "Composite risk score across all sensors"),
    ("GET",  "/api/pipeline/topology",            "Full pipeline topology: nodes + edges for the map"),
    ("GET",  "/api/pipeline/nodes",               "List all 9 pipeline nodes"),
    ("POST", "/api/reports/generate",             "Generate a report. Body: {report_type, hours}"),
    ("GET",  "/api/reports",                      "List all generated reports"),
    ("GET",  "/api/reports/sensor-health",        "24-hour health metrics for all sensors"),
    ("WS",   "/ws/pipeline",                      "WebSocket — streams ALL sensor readings and alerts"),
]

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
t.rows[0].cells[0].text = "Method"
t.rows[0].cells[1].text = "Endpoint"
t.rows[0].cells[2].text = "Description"
shade_row(t.rows[0], 'BDD7EE')
for i in range(3):
    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    t.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(9)
for method, path, desc in endpoints:
    row = t.add_row()
    row.cells[0].text = method
    row.cells[1].text = path
    row.cells[2].text = desc
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[1].paragraphs[0].runs[0].font.name = "Courier New"
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(8)
    row.cells[2].paragraphs[0].runs[0].font.size = Pt(9)
    if method == "POST":
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xAA, 0x55, 0x00)
    elif method == "WS":
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00, 0x88, 0x44)

doc.add_page_break()

# ── SECTION 11: WEBSOCKET MESSAGE FORMAT ─────────
set_heading(doc, "11.  WEBSOCKET MESSAGES — What the App Receives", 1, (0x00, 0x55, 0x88))

add_para(doc,
    "Every 2 seconds, the backend sends a sensor_reading message to the app. "
    "When a leak is detected, it sends an alert message immediately (real-time).", size=10)

add_para(doc, "Sensor Reading Message:", bold=True, size=10)
add_code(doc,
'{\n'
'  "type": "sensor_reading",\n'
'  "payload": {\n'
'    "sensor_id": "PS-001",          // Which sensor sent this\n'
'    "value": 103.5,                  // The actual reading (PSI, L/min, dB, or C)\n'
'    "unit": "PSI",                   // Unit of measurement\n'
'    "timestamp": "2026-03-30T10:15:23.123Z",\n'
'    "is_anomaly": false,             // Z-score > 3.0 threshold?\n'
'    "confidence_score": 0.92,        // 0.0-1.0 (how confident the detection is)\n'
'    "leak_probability": 0.15,        // 0.0-1.0 (likelihood this is a leak)\n'
'    "agent_state": "NORMAL"          // "NORMAL" or "LEAK"\n'
'  }\n'
'}')

doc.add_paragraph()
add_para(doc, "Alert Message (sent immediately when anomaly detected):", bold=True, size=10)
add_code(doc,
'{\n'
'  "type": "alert",\n'
'  "payload": {\n'
'    "id": 42,\n'
'    "sensor_id": "PS-001",\n'
'    "alert_type": "ABNORMAL_PRESSURE_DROP",\n'
'    "severity": "CRITICAL",          // CRITICAL / WARNING / INFO\n'
'    "message": "PS-001 dropped 27 PSI in 8 seconds. Leak probability 94%.",\n'
'    "timestamp": "2026-03-30T10:15:40.000Z",\n'
'    "value_at_trigger": 76.2,        // The reading that triggered the alert\n'
'    "location_label": "Sector 7 — Km 12.4",\n'
'    "acknowledged": false,\n'
'    "resolved": false\n'
'  }\n'
'}')

doc.add_page_break()

# ── SECTION 12: ABBREVIATIONS ─────────────────────
set_heading(doc, "12.  ABBREVIATIONS AND TERMS GLOSSARY", 1, (0x00, 0x55, 0x88))

terms = [
    ("PSI", "Pounds per Square Inch — unit of pressure. 1 PSI ≈ 6.895 kPa. Normal pipeline pressure: 85-115 PSI."),
    ("L/min", "Litres per Minute — unit of flow rate. Normal flow: 180-220 L/min."),
    ("dB", "Decibels — unit of sound level. Normal acoustic background: 28-55 dB. Leak spike: up to 90 dB."),
    ("°C", "Degrees Celsius — unit of temperature. Normal ground temperature: 15-28°C."),
    ("Z-score", "Statistical measure of how unusual a value is. Z > 3.0 = anomaly. Z = (value - mean) / standard_deviation."),
    ("Anomaly", "A reading that is statistically unusual — Z-score above 3.0. Not all anomalies are leaks, but all leaks cause anomalies."),
    ("KPI", "Key Performance Indicator — the 4 big number cards on the Dashboard showing live sensor values."),
    ("WebSocket", "A persistent two-way connection between phone and backend. Unlike HTTP (request/response), WebSocket allows the server to PUSH data to the phone at any time."),
    ("FastAPI", "Python web framework for building APIs. Async = can handle many connections simultaneously without blocking."),
    ("SQLite", "A file-based database stored in ecoguard.db. No separate database server needed."),
    ("SQLAlchemy", "Python ORM (Object Relational Mapper) — lets Python code work with database tables as Python objects."),
    ("Expo", "Platform for building React Native apps. Expo Go = phone app for testing without building a full native app."),
    ("React Context", "A way to share data (like theme colors) across many components without passing props manually."),
    ("Agent", "An autonomous software process that runs independently. In EcoGuard, each sensor has its own agent running as an asyncio background task."),
    ("Asyncio", "Python's framework for running multiple tasks concurrently. Allows 4 sensor agents + web server + WebSocket all in one process."),
    ("CORS", "Cross-Origin Resource Sharing — security setting that allows the mobile app to call the backend API across different domains/ports."),
    ("ngrok", "Tool that creates a public HTTPS URL tunneling to your local server. Used when the phone needs to access the backend via the internet instead of local WiFi."),
    ("Nominal", "Normal operating range. When a sensor reads within its nominal range, status = NOMINAL (green)."),
    ("SDK", "Software Development Kit — Expo SDK 54 is the version of Expo's tools and libraries used."),
    ("Composite Risk", "The combined leak risk score calculated from all 4 sensors with different weights: Pressure 35%, Flow 30%, Acoustic 25%, IR 10%."),
    ("Debounce", "A timer that prevents duplicate events. The 30-second alert debounce means one sensor can only create one alert per 30 seconds."),
    ("Piezoelectric", "A material that generates electrical signal from mechanical stress. Piezoelectric sensors are used in pressure transducers and acoustic emission sensors."),
]

t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
t.rows[0].cells[0].text = "Term / Abbreviation"
t.rows[0].cells[1].text = "Meaning"
shade_row(t.rows[0], 'BDD7EE')
for i in range(2):
    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    t.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(9)
for term, meaning in terms:
    row = t.add_row()
    row.cells[0].text = term
    row.cells[1].text = meaning
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)

doc.add_page_break()

# ── SECTION 13: PRESENTATION TALKING POINTS ───────
set_heading(doc, "13.  WHAT TO SAY WHEN PRESENTING", 1, (0x00, 0x55, 0x88))

add_para(doc, "Opening Statement", bold=True, size=12)
add_para(doc,
    '"EcoGuard is an industrial pipeline leakage detection system that monitors a 48 km pipeline '
    'using 4 autonomous sensor agents. It uses statistical anomaly detection and real-time '
    'WebSocket streaming to detect leaks within seconds and alert operators on their mobile device."',
    size=10, italic=True)

doc.add_paragraph()
for topic, point in [
    ("If asked about the backend:",
     "\"The backend is a FastAPI Python server with 4 autonomous agents running as asyncio background tasks. "
     "Each agent simulates a real sensor with Gaussian noise and sinusoidal drift, then uses Z-score "
     "statistical analysis to detect anomalies. Readings are persisted to SQLite via SQLAlchemy ORM.\""),
    ("If asked about real-time:",
     "\"The system uses WebSocket protocol — not HTTP polling. The backend pushes data to the mobile app "
     "every 2 seconds without the app having to request it. Alerts are pushed immediately when detected, "
     "giving true real-time response in under 1 second.\""),
    ("If asked about the detection algorithm:",
     "\"Each sensor computes a Z-score: (current_value minus rolling_mean) divided by standard_deviation. "
     "If Z exceeds 3.0, it's an anomaly. The four sensor probabilities are then combined with weights: "
     "35% pressure, 30% flow, 25% acoustic, 10% infrared to give a composite risk score.\""),
    ("If asked about the theme system:",
     "\"The app uses React Context for global theme management. A ThemeProvider wraps the entire app, "
     "and every component calls useTheme() to get dynamic colors. Toggling dark/light/system instantly "
     "re-renders all components with the new color set — including backgrounds, text, borders, and accents.\""),
    ("If asked about scalability:",
     "\"The architecture is designed for production scale. The database uses SQLAlchemy ORM so switching "
     "from SQLite to PostgreSQL is one config line change. The agent manager can add more sensor types. "
     "WebSocket connections are managed centrally and support multiple concurrent mobile clients.\""),
]:
    set_heading(doc, topic, 3, (0x00, 0x55, 0x88))
    add_para(doc, point, size=10, italic=True)
    doc.add_paragraph()

# ── FOOTER ────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f"EcoGuard Technologies  ·  Pipeline Integrity Platform  ·  "
              f"Generated {datetime.date.today().strftime('%B %d, %Y')}")
r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
r.italic = True

# ── SAVE ─────────────────────────────────────────
out_path = r"c:\Users\hp\OneDrive\Desktop\projects\ECOGUARD\EcoGuard_Complete_Documentation.docx"
doc.save(out_path)
print(f"Documentation saved to:\n  {out_path}")
